"""
Kearney Document Engine - Brand-Compliant Word Document Generation

Generates Word documents that conform to Kearney Design System (KDS) standards.
This module wraps python-docx to enforce brand compliance automatically.

Usage:
    from core.document_engine import KDSDocument

    doc = KDSDocument()
    doc.add_cover_page(title='Q3 Analysis', client='Acme Corp')
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('Key findings from our analysis...')
    doc.save('exports/report.docx')
"""

from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Tuple
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Kearney Brand Colors (RGB)
KEARNEY_PURPLE = RGBColor(120, 35, 220) if DOCX_AVAILABLE else None
TEXT_DARK = RGBColor(51, 51, 51) if DOCX_AVAILABLE else None
TEXT_MEDIUM = RGBColor(102, 102, 102) if DOCX_AVAILABLE else None
WHITE = RGBColor(255, 255, 255) if DOCX_AVAILABLE else None


class KDSDocument:
    """
    Kearney Design System compliant document generator.

    Enforces brand standards automatically:
    - Kearney Purple accents
    - Inter/Arial typography
    - Consistent formatting
    - No green colors
    - No emojis
    """

    def __init__(self):
        """Initialize a new KDS-compliant document."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for document generation. "
                "Install with: pip install python-docx"
            )

        self.document = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Set up KDS-compliant styles."""
        # Configure default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'  # Inter fallback
        font.size = Pt(11)
        font.color.rgb = TEXT_DARK

        # Configure paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15

        # Configure heading styles
        self._setup_heading_styles()

    def _setup_heading_styles(self):
        """Configure heading styles with Kearney brand."""
        heading_configs = [
            ('Heading 1', 18, True, KEARNEY_PURPLE),
            ('Heading 2', 14, True, KEARNEY_PURPLE),
            ('Heading 3', 12, True, TEXT_DARK),
            ('Heading 4', 11, True, TEXT_DARK),
        ]

        for style_name, size, bold, color in heading_configs:
            try:
                style = self.document.styles[style_name]
                font = style.font
                font.name = 'Arial'
                font.size = Pt(size)
                font.bold = bold
                font.color.rgb = color
            except KeyError:
                # Style doesn't exist, skip
                pass

    def add_cover_page(
        self,
        title: str,
        client: str = "",
        date: str = "",
        author: str = "",
        confidential: bool = True
    ) -> None:
        """
        Add a cover page with project metadata.

        Args:
            title: Document title
            client: Client name
            date: Date string (defaults to today)
            author: Author name
            confidential: Whether to mark as confidential
        """
        if not date:
            date = datetime.now().strftime("%B %Y")

        # Add spacing at top
        for _ in range(4):
            self.document.add_paragraph()

        # Title
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = KEARNEY_PURPLE
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some spacing
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Client and date
        if client:
            client_para = self.document.add_paragraph()
            client_run = client_para.add_run(f"Prepared for: {client}")
            client_run.font.size = Pt(14)
            client_run.font.color.rgb = TEXT_MEDIUM
            client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = self.document.add_paragraph()
        date_run = date_para.add_run(date)
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = TEXT_MEDIUM
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if author:
            author_para = self.document.add_paragraph()
            author_run = author_para.add_run(f"Author: {author}")
            author_run.font.size = Pt(11)
            author_run.font.color.rgb = TEXT_MEDIUM
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Confidentiality notice at bottom
        if confidential:
            for _ in range(8):
                self.document.add_paragraph()
            conf_para = self.document.add_paragraph()
            conf_run = conf_para.add_run("CONFIDENTIAL")
            conf_run.font.size = Pt(10)
            conf_run.font.italic = True
            conf_run.font.color.rgb = TEXT_MEDIUM
            conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break after cover
        self.document.add_page_break()

    def add_heading(self, text: str, level: int = 1) -> None:
        """
        Add a heading with KDS styling.

        Args:
            text: Heading text
            level: Heading level (1-4)
        """
        if level < 1:
            level = 1
        if level > 4:
            level = 4

        self.document.add_heading(text, level=level)

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        style: str = None
    ) -> None:
        """
        Add a paragraph with optional formatting.

        Args:
            text: Paragraph text
            bold: Whether to bold the text
            italic: Whether to italicize the text
            style: Optional style name to apply
        """
        para = self.document.add_paragraph(style=style)
        run = para.add_run(text)
        run.font.bold = bold
        run.font.italic = italic

    def add_bullet_list(self, items: List[str]) -> None:
        """
        Add a bulleted list.

        Args:
            items: List of bullet point texts
        """
        for item in items:
            para = self.document.add_paragraph(item, style='List Bullet')

    def add_numbered_list(self, items: List[str]) -> None:
        """
        Add a numbered list.

        Args:
            items: List of items
        """
        for item in items:
            para = self.document.add_paragraph(item, style='List Number')

    def add_table(
        self,
        data: List[List[Any]],
        headers: List[str],
        caption: str = None
    ) -> None:
        """
        Add a KDS-styled table.

        Args:
            data: List of rows (each row is a list of values)
            headers: Column headers
            caption: Optional caption below table
        """
        # Create table with header row
        num_cols = len(headers)
        table = self.document.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add headers with purple background
        header_row = table.rows[0]
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = str(header_text)

            # Style header cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = WHITE
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Set background color (requires shading)
            self._set_cell_shading(cell, "7823DC")

        # Add data rows
        for row_data in data:
            row = table.add_row()
            for idx, value in enumerate(row_data):
                cell = row.cells[idx]
                cell.text = str(value)

        # Add caption if provided
        if caption:
            caption_para = self.document.add_paragraph()
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            caption_run.font.color.rgb = TEXT_MEDIUM

    def _set_cell_shading(self, cell, color_hex: str):
        """Set cell background color."""
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    def add_chart(
        self,
        image_path: Union[str, Path],
        caption: str = None,
        width_inches: float = 5.5
    ) -> None:
        """
        Add a chart image from file.

        Args:
            image_path: Path to chart image (PNG, JPG)
            caption: Optional caption below chart
            width_inches: Width in inches (height scales proportionally)
        """
        image_path = Path(image_path)
        if not image_path.exists():
            raise FileNotFoundError(f"Chart image not found: {image_path}")

        self.document.add_picture(str(image_path), width=Inches(width_inches))

        # Center the image
        last_paragraph = self.document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add caption if provided
        if caption:
            caption_para = self.document.add_paragraph()
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            caption_run.font.color.rgb = TEXT_MEDIUM
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_page_break(self) -> None:
        """Add a page break."""
        self.document.add_page_break()

    def add_section_break(self) -> None:
        """Add a section break (starts new section)."""
        from docx.enum.section import WD_ORIENT
        self.document.add_section()

    def add_executive_summary(self, points: List[str]) -> None:
        """
        Add an executive summary section.

        Args:
            points: Key points to include
        """
        self.add_heading("Executive Summary", level=1)

        for point in points:
            para = self.document.add_paragraph()
            run = para.add_run(f"- {point}")
            run.font.size = Pt(11)

    def add_recommendations(self, recommendations: List[str]) -> None:
        """
        Add a recommendations section.

        Args:
            recommendations: List of recommendations
        """
        self.add_heading("Recommendations", level=1)
        self.add_numbered_list(recommendations)

    def add_next_steps(self, steps: List[Tuple[str, str, str]]) -> None:
        """
        Add a next steps section with action items.

        Args:
            steps: List of (action, owner, due_date) tuples
        """
        self.add_heading("Next Steps", level=1)

        headers = ["Action", "Owner", "Due Date"]
        data = [[action, owner, due] for action, owner, due in steps]
        self.add_table(data, headers)

    def save(self, path: Union[str, Path]) -> Path:
        """
        Save the document to file.

        Args:
            path: Output file path

        Returns:
            Path to the saved file
        """
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(path)
        return path


def create_memo(
    to: str,
    from_sender: str,
    subject: str,
    situation: str,
    analysis: str,
    recommendation: str,
    next_steps: List[str] = None,
    output_path: Union[str, Path] = None,
    date: str = None
) -> Path:
    """
    Convenience function to create a memo document.

    Args:
        to: Recipient
        from_sender: Sender
        subject: Memo subject
        situation: Situation/background text
        analysis: Analysis text
        recommendation: Recommendation text
        next_steps: Optional list of next steps
        output_path: Where to save the file
        date: Date (defaults to today)

    Returns:
        Path to the saved file
    """
    doc = KDSDocument()

    if not date:
        date = datetime.now().strftime("%Y-%m-%d")

    # Memo header block
    doc.add_paragraph(f"TO: {to}", bold=True)
    doc.add_paragraph(f"FROM: {from_sender}", bold=True)
    doc.add_paragraph(f"DATE: {date}", bold=True)
    doc.add_paragraph(f"RE: {subject}", bold=True)

    # Horizontal line (using underscores)
    doc.add_paragraph("_" * 60)

    # Sections
    doc.add_heading("Situation", level=2)
    doc.add_paragraph(situation)

    doc.add_heading("Analysis", level=2)
    doc.add_paragraph(analysis)

    doc.add_heading("Recommendation", level=2)
    doc.add_paragraph(recommendation)

    if next_steps:
        doc.add_heading("Next Steps", level=2)
        doc.add_bullet_list(next_steps)

    if output_path is None:
        safe_subject = subject.lower().replace(' ', '_')[:30]
        output_path = f"outputs/memo_{safe_subject}.docx"

    return doc.save(output_path)


def create_report(
    title: str,
    client: str,
    executive_summary: List[str],
    findings: List[Tuple[str, str]],
    recommendations: List[str],
    output_path: Union[str, Path] = None,
    author: str = ""
) -> Path:
    """
    Convenience function to create a standard report.

    Args:
        title: Report title
        client: Client name
        executive_summary: List of key points
        findings: List of (finding_title, finding_text) tuples
        recommendations: List of recommendations
        output_path: Where to save the file
        author: Author name

    Returns:
        Path to the saved file
    """
    doc = KDSDocument()

    doc.add_cover_page(title=title, client=client, author=author)

    doc.add_executive_summary(executive_summary)

    doc.add_heading("Findings", level=1)
    for finding_title, finding_text in findings:
        doc.add_heading(finding_title, level=2)
        doc.add_paragraph(finding_text)

    doc.add_recommendations(recommendations)

    if output_path is None:
        safe_title = title.lower().replace(' ', '_')[:30]
        output_path = f"exports/{safe_title}.docx"

    return doc.save(output_path)
